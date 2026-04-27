#!/usr/bin/env python3
import sys, os, csv
from docx import Document

def csv_to_word(dirs, out='output.docx'):
    doc = Document()
    for d in dirs:
        if not os.path.isdir(d):
            continue
        # get CSV files (first 2 if you strictly want only two)
        csv_files = sorted([f for f in os.listdir(d) if f.lower().endswith('.csv')])[:]
        for fname in csv_files:
            path = os.path.join(d, fname)
            with open(path, newline='', encoding='utf-8') as f:
                reader = list(csv.reader(f))
                if not reader:
                    continue
                # Add caption paragraph with style 'Caption' (Word caption)
                para = doc.add_paragraph(fname)
                try:
                    para.style = 'Caption'
                except Exception:
                    pass
                # Build table
                maxcols = max(len(r) for r in reader)
                table = doc.add_table(rows=len(reader), cols=maxcols)
                table.style = 'Table Grid'
                for i, row in enumerate(reader):
                    for j in range(maxcols):
                        text = row[j] if j < len(row) else ''
                        table.cell(i, j).text = str(text)
                doc.add_paragraph()  # spacer
    doc.save(out)

if __name__ == '__main__':
    if len(sys.argv) != 4:
        print("Usage: python3 csvs_to_word.py dir1 dir2 dir3")
        sys.exit(1)
    csv_to_word(sys.argv[1:4])